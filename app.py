import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO

# Configuração da Página
st.set_page_config(page_title="Portal Perito RS", layout="wide")

# Recuperação das Chaves (Secrets)
api_key = st.secrets.get("GOOGLE_API_KEY")

if api_key:
    genai.configure(api_key=api_key)
else:
    st.error("Erro: GOOGLE_API_KEY não configurada.")

st.title("🚀 Portal Perito RS")

# Diagnóstico Corrigido (Sem o comando que causa erro)
if st.button("🔍 Rodar Diagnóstico"):
    try:
        # Apenas lista os nomes dos modelos, comando que funciona em todas as versões
        modelos = [m.name for m in genai.list_models()]
        st.success("Conectado com sucesso ao Google AI!")
        st.write("Modelos disponíveis:", modelos)
    except Exception as e:
        st.error(f"Erro na conexão: {e}")

st.divider()

materia = st.text_input("Qual a matéria do planejamento?")
ano = st.selectbox("Para qual ano?", ["1º Ano", "2º Ano", "3º Ano", "4º Ano", "5º Ano"])

if st.button("Gerar Planejamento"):
    if materia:
        with st.spinner("Gerando plano em Arial 12..."):
            try:
                # Usa o modelo Gemini 1.5 Flash diretamente
                model = genai.GenerativeModel('gemini-1.5-flash')
                response = model.generate_content(f"Crie um plano de aula completo para {materia}, {ano}, seguindo a BNCC.")
                
                texto_gerado = response.text
                st.subheader("Resultado:")
                st.write(texto_gerado)
                
                # Criando o arquivo Word formatado
                doc = Document()
                style = doc.styles['Normal']
                style.font.name = 'Arial'
                style.font.size = 12
                
                doc.add_heading(f'Plano de Aula: {materia} - {ano}', 0)
                doc.add_paragraph(texto_gerado)
                
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                st.download_button(
                    label="📥 Baixar Plano em Word",
                    data=buffer,
                    file_name=f"Plano_{materia}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"Erro ao gerar: {e}")
